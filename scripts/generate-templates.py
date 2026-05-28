#!/usr/bin/env python3
"""generate-templates.py — produces all docx templates per style-config.json.

Each template includes a heading, a body paragraph, a statute-quote block,
and a signature block, all formatted per the law-firm style.
"""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
STYLE_CONFIG_PATH = ROOT / "templates" / "_style" / "style-config.json"
TEMPLATE_DIR = ROOT / "templates"


def _set_run_cn_font(run, name: str, size_pt: float, bold: bool = False):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")


def _set_page(doc, cfg):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(cfg["page"]["margin_top_cm"])
    section.bottom_margin = Cm(cfg["page"]["margin_bottom_cm"])
    section.left_margin = Cm(cfg["page"]["margin_left_cm"])
    section.right_margin = Cm(cfg["page"]["margin_right_cm"])


def _add_para(doc, text, font, size, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
              first_line_indent_chars=0, line_spacing_pt=28, space_before_pt=0,
              space_after_pt=0):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line_spacing_pt)
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)
    if first_line_indent_chars:
        pf.first_line_indent = Pt(size * first_line_indent_chars)
    run = p.add_run(text)
    _set_run_cn_font(run, font, size, bold)
    return p


def _build(doc_title, body_sections, signature_lines, cfg):
    doc = Document()
    _set_page(doc, cfg)

    # Title (二号 方正小标宋简体, centered, no indent)
    _add_para(doc, doc_title,
              cfg["fonts"]["title_cn"], cfg["sizes_pt"]["title"],
              align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent_chars=0,
              line_spacing_pt=cfg["line_spacing_pt"],
              space_before_pt=12, space_after_pt=18)

    for heading, paragraphs, statute_quote in body_sections:
        # H1
        if heading:
            _add_para(doc, heading,
                      cfg["fonts"]["h1_cn"], cfg["sizes_pt"]["h1"],
                      bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
                      first_line_indent_chars=0,
                      line_spacing_pt=cfg["line_spacing_pt"],
                      space_before_pt=12, space_after_pt=6)
        for para in paragraphs:
            _add_para(doc, para,
                      cfg["fonts"]["body_cn"], cfg["sizes_pt"]["body"],
                      first_line_indent_chars=cfg["first_line_indent_chars"],
                      line_spacing_pt=cfg["line_spacing_pt"])
        if statute_quote:
            _add_para(doc, statute_quote,
                      cfg["fonts"]["statute_quote_cn"], cfg["sizes_pt"]["statute_quote"],
                      align=WD_ALIGN_PARAGRAPH.LEFT,
                      first_line_indent_chars=cfg["first_line_indent_chars"],
                      line_spacing_pt=cfg["line_spacing_pt"])

    # Signature block (right-aligned, KaiTi)
    if signature_lines:
        for line in signature_lines:
            _add_para(doc, line,
                      cfg["fonts"]["signature_cn"], cfg["sizes_pt"]["body"],
                      align=WD_ALIGN_PARAGRAPH.RIGHT,
                      first_line_indent_chars=0,
                      line_spacing_pt=cfg["line_spacing_pt"],
                      space_before_pt=18)
    return doc


TEMPLATES = [
    ("legal-memo-cn.docx", "法律备忘录",
     [
        ("一、问题",
         ["本备忘录就贵司询问的问题事项进行分析。",
          "在此填入问题概要。"], None),
        ("二、简要结论",
         ["简要结论应当回应问题，不超过三段。"], None),
        ("三、事实",
         ["在此列出与法律分析相关的事实，按时间顺序排列。"], None),
        ("四、法律分析",
         ["在此对每一争议点逐一进行法律分析，引用具体条文与案例。"],
         "《民法典》第五百零九条　当事人应当按照约定全面履行自己的义务。"),
        ("五、结论与建议",
         ["综上，本所建议如下："], None),
     ],
     ["签发律师：__________", "出具日期：__________"],
     ),

    ("legal-opinion-cn.docx", "法律意见书",
     [
        ("一、当事人与受托范围",
         ["在此说明委托人、受托范围与本意见的依据。"], None),
        ("二、假设事项",
         ["本意见依赖以下假设事项："], None),
        ("三、限定事项",
         ["本意见受以下限定事项的约束："], None),
        ("四、法律意见",
         ["基于上述假设与限定，本所谨提供以下法律意见："],
         "《公司法》第十四条　公司可以设立分公司。"),
     ],
     ["律师事务所（盖章）：__________",
      "经办律师：__________ / __________",
      "出具日期：__________"]),

    ("civil-complaint-cn.docx", "民事起诉状",
     [
        ("原告",
         ["姓名 / 名称：__________",
          "性别 / 注册号：__________",
          "住所 / 注册地：__________",
          "联系方式：__________"], None),
        ("被告",
         ["姓名 / 名称：__________",
          "住所 / 注册地：__________"], None),
        ("诉讼请求",
         ["1. ____________________；",
          "2. ____________________；",
          "3. 由被告承担本案诉讼费用。"], None),
        ("事实与理由",
         ["在此叙述事实与理由，按时间顺序展开。",
          "据此，依据《民事诉讼法》及相关规定，特向贵院提起诉讼，请依法判处。"],
         "《民法典》第五百七十七条　当事人一方不履行合同义务或者履行合同义务不符合约定的，应当承担继续履行、采取补救措施或者赔偿损失等违约责任。"),
     ],
     ["此致", "____________人民法院",
      "起诉人（签名 / 盖章）：__________",
      "____________年____月____日"]),

    ("civil-defence-cn.docx", "民事答辩状",
     [
        ("答辩人", ["姓名 / 名称：__________", "住所：__________"], None),
        ("被答辩人", ["姓名 / 名称：__________"], None),
        ("答辩意见",
         ["原告诉讼请求缺乏事实与法律依据，理由如下：",
          "一、____________________；",
          "二、____________________；",
          "三、____________________。"], None),
        ("结论",
         ["综上，请贵院依法驳回原告的全部诉讼请求。"], None),
     ],
     ["此致", "____________人民法院",
      "答辩人（签名 / 盖章）：__________",
      "____________年____月____日"]),

    ("appellate-brief-cn.docx", "民事上诉状",
     [
        ("上诉人", ["姓名 / 名称：__________"], None),
        ("被上诉人", ["姓名 / 名称：__________"], None),
        ("原审情况", ["原审法院：__________", "原审案号：__________", "原审判决主要内容：__________"], None),
        ("上诉请求",
         ["1. 撤销 ____________________；",
          "2. 改判 ____________________；",
          "3. 一审、二审诉讼费用由被上诉人承担。"], None),
        ("事实与理由",
         ["一、原判决认定事实不清，理由如下：____________________。",
          "二、原判决适用法律错误，理由如下：____________________。"], None),
     ],
     ["此致", "____________人民法院",
      "上诉人（签名 / 盖章）：__________",
      "____________年____月____日"]),

    ("evidence-index-cn.docx", "证据清单",
     [
        ("一、证据目录",
         ["1. 证据名称：__________；证明目的：__________；证据来源：__________。",
          "2. 证据名称：__________；证明目的：__________；证据来源：__________。",
          "3. 证据名称：__________；证明目的：__________；证据来源：__________。"], None),
        ("二、说明",
         ["上述证据均为客观真实，与本案具有关联性，特此提交。"], None),
     ],
     ["提交人（签名 / 盖章）：__________",
      "____________年____月____日"]),

    ("settlement-agreement-cn.docx", "和解协议",
     [
        ("甲方", ["名称：__________", "住所：__________"], None),
        ("乙方", ["名称：__________", "住所：__________"], None),
        ("鉴于条款",
         ["鉴于双方就 ____________________ 事项存在争议，本着平等、自愿、公平的原则，双方就争议事项达成和解，签订本协议。"], None),
        ("和解条款",
         ["一、付款安排：__________；",
          "二、释放范围：__________；",
          "三、保密条款：__________；",
          "四、违约责任：__________；",
          "五、争议解决：__________。"], None),
     ],
     ["甲方（签名 / 盖章）：__________",
      "乙方（签名 / 盖章）：__________",
      "签订日期：__________"]),

    ("sale-contract-cn.docx", "货物买卖合同",
     [
        ("一、合同主体", ["卖方：__________", "买方：__________"], None),
        ("二、标的", ["货物名称、规格、数量：__________"], None),
        ("三、价款与支付", ["价款金额：__________；支付方式：__________；支付时间：__________"], None),
        ("四、交付", ["交付地点：__________；交付方式：__________；风险转移：__________"], None),
        ("五、质量与验收", ["质量标准：__________；验收期：__________"], None),
        ("六、违约责任", ["如一方违反本合同，应承担继续履行、补救或赔偿损失等违约责任。"],
         "《民法典》第六百一十条 出卖人交付的标的物不符合质量要求的，买受人可以依法解除合同或者请求承担违约责任。"),
        ("七、争议解决", ["双方因本合同发生争议的，应友好协商解决；协商不成，提交 ____________ 仲裁委员会仲裁。"], None),
     ],
     ["卖方（盖章）：__________", "买方（盖章）：__________", "签订日期：__________"]),

    ("services-contract-cn.docx", "服务合同",
     [
        ("一、合同主体", ["委托方：__________", "服务方：__________"], None),
        ("二、服务内容", ["服务范围：__________；交付物：__________；验收标准：__________"], None),
        ("三、服务费用", ["费用金额：__________；支付方式：__________"], None),
        ("四、知识产权", ["工作成果的知识产权归属：__________"], None),
        ("五、保密义务", ["服务方应对在履行本合同过程中知悉的委托方商业秘密承担保密义务。"], None),
        ("六、违约与解除", ["违约责任及合同解除条件如下：__________"], None),
     ],
     ["委托方（盖章）：__________", "服务方（盖章）：__________", "签订日期：__________"]),

    ("nda-mutual-cn.docx", "双向保密协议",
     [
        ("一、合同主体", ["甲方：__________", "乙方：__________"], None),
        ("二、保密信息", ["保密信息的范围：__________"], None),
        ("三、保密义务", ["双方应对保密信息承担同等的保密义务。"], None),
        ("四、保密期限", ["自本协议签订之日起 ____ 年。"], None),
        ("五、违约责任", ["违反本协议的一方应向守约方支付违约金 ____________ 元，并赔偿实际损失。"], None),
     ],
     ["甲方（盖章）：__________", "乙方（盖章）：__________", "签订日期：__________"]),

    ("labor-contract-cn.docx", "劳动合同",
     [
        ("一、合同主体", ["用人单位：__________", "劳动者：__________"], None),
        ("二、合同期限", ["合同类型：固定期限 / 无固定期限 / 以完成一定工作任务为期限。", "起止日期：__________"], None),
        ("三、工作内容与地点", ["岗位：__________", "工作地点：__________", "工作时间：__________"], None),
        ("四、劳动报酬", ["月工资标准：__________", "发放方式：__________"], None),
        ("五、社会保险", ["用人单位依法为劳动者缴纳社会保险及住房公积金。"], None),
        ("六、保密与竞业限制", ["劳动者应对用人单位的商业秘密负保密义务。竞业限制条款详见附件。"], None),
        ("七、合同解除与终止", ["依据《劳动合同法》及双方约定执行。"], None),
     ],
     ["用人单位（盖章）：__________", "劳动者（签名）：__________", "签订日期：__________"]),

    ("termination-letter-cn.docx", "劳动合同解除通知书",
     [
        ("",
         ["__________先生 / 女士："], None),
        ("",
         ["根据双方于 ________ 签订的《劳动合同》及《劳动合同法》第____条之规定，本公司决定自 ____________ 日起解除与您的劳动合同。",
          "经济补偿金为人民币 ____________ 元（大写：__________），将依法支付至您指定的账户。",
          "请您于 ________ 前办理工作交接，并按公司规定办理离职手续。",
          "感谢您在职期间为本公司所作的贡献。"], None),
     ],
     ["__________有限公司（盖章）：__________", "签发日期：__________"]),

    ("non-compete-clause-cn.docx", "竞业限制协议",
     [
        ("一、协议主体", ["用人单位：__________", "劳动者：__________"], None),
        ("二、竞业限制范围",
         ["竞业限制的地域：__________", "竞业限制的行业：__________", "禁止从事的工作类型：__________"], None),
        ("三、竞业限制期限", ["自劳动者离职之日起 ____ 年（最长不超过两年）。"], None),
        ("四、经济补偿",
         ["竞业限制期间，用人单位按月向劳动者支付经济补偿，标准为劳动者离职前 12 个月平均工资的 ____ %（不低于法律规定的最低标准）。"], None),
        ("五、违约责任", ["劳动者违反本协议的，应向用人单位支付违约金 ____________ 元，并赔偿实际损失。"], None),
     ],
     ["用人单位（盖章）：__________", "劳动者（签名）：__________", "签订日期：__________"]),

    ("articles-of-association-cn.docx", "公司章程",
     [
        ("第一章 总则", ["第一条 为规范公司的组织和行为，依据《公司法》及有关规定，制定本章程。"], None),
        ("第二章 公司名称与住所",
         ["第二条 公司名称：__________", "第三条 公司住所：__________"], None),
        ("第三章 经营范围", ["第四条 经营范围：__________"], None),
        ("第四章 注册资本与股东", ["第五条 注册资本：人民币 ____________ 元", "第六条 股东及其出资如下：__________"], None),
        ("第五章 公司机构", ["第七条 公司设股东会、董事会、监事会；其职责如下：__________"], None),
        ("第六章 财务与利润分配", ["第八条 公司利润按法律规定与本章程分配。"], None),
        ("第七章 终止与清算", ["第九条 公司有下列情形之一的，应当解散：__________"], None),
        ("第八章 附则", ["第十条 本章程自公司登记机关核准之日起生效。"], None),
     ],
     ["全体股东签字 / 盖章：__________", "签订日期：__________"]),

    ("board-resolution-cn.docx", "董事会决议",
     [
        ("",
         ["__________有限公司董事会于 __________ 在 __________ 召开。",
          "出席董事 __ 人；缺席董事 __ 人；列席人员 __ 人；会议主持人：__________。",
          "经审议，董事会一致通过 / 多数通过如下决议："], None),
        ("决议事项",
         ["一、____________________；",
          "二、____________________；",
          "三、____________________。"], None),
        ("",
         ["本决议自全体董事签字之日起生效。"], None),
     ],
     ["全体董事签字：__________", "签署日期：__________"]),

    ("shareholder-resolution-cn.docx", "股东会决议",
     [
        ("",
         ["__________有限公司股东会于 __________ 在 __________ 召开。",
          "出席股东持有公司表决权 ____ %；会议主持人：__________。",
          "经审议，股东会以 ____ %表决权通过下列决议："], None),
        ("决议事项",
         ["一、____________________；",
          "二、____________________；",
          "三、____________________。"], None),
     ],
     ["全体股东签字 / 盖章：__________", "签署日期：__________"]),

    ("regulator-response-cn.docx", "监管问询答复",
     [
        ("",
         ["致：____________________（监管机关）"], None),
        ("",
         ["本公司收到贵机关 ________ 号《问询函》（以下简称该问询函）。本公司高度重视，组织相关部门进行核查，现就问询事项答复如下："], None),
        ("一、关于事项一",
         ["问询内容摘要：____________________",
          "本公司答复：____________________"], None),
        ("二、关于事项二",
         ["问询内容摘要：____________________",
          "本公司答复：____________________"], None),
        ("",
         ["本公司承诺，本答复所述事项真实、准确、完整，不存在虚假记载、误导性陈述或重大遗漏。"], None),
     ],
     ["__________有限公司（盖章）：__________", "答复日期：__________"]),

    ("legal-due-diligence-report-cn.docx", "法律尽职调查报告",
     [
        ("一、调查范围与方法",
         ["本所受 __________ 委托，对 __________ 进行法律尽职调查。调查范围、方法及限定事项如下：__________"], None),
        ("二、公司基本情况",
         ["公司名称：__________；统一社会信用代码：__________；注册资本：__________；股东与持股比例：__________；公司治理结构：__________"], None),
        ("三、经营与资质", ["主要业务：__________；经营许可与资质：__________"], None),
        ("四、重大合同", ["重大合同清单及风险点：__________"], None),
        ("五、劳动用工", ["员工人数：__________；社保住房公积金：__________；劳动争议：__________"], None),
        ("六、知识产权", ["商标、专利、著作权清单及权利状态：__________"], None),
        ("七、不动产", ["土地使用权与房屋产权情况：__________"], None),
        ("八、税务与财务", ["近三年税务情况：__________"], None),
        ("九、诉讼仲裁与行政处罚", ["未决诉讼仲裁：__________；行政处罚：__________"], None),
        ("十、结论与建议", ["综上，本所认为：__________；建议：__________"], None),
     ],
     ["律师事务所（盖章）：__________",
      "经办律师：__________ / __________",
      "出具日期：__________"]),

    ("demand-letter-cn.docx", "律师函",
     [
        ("",
         ["致：__________"], None),
        ("",
         ["本所接受 __________（以下简称委托人）的委托，就贵方与委托人之间 __________ 事宜，特此函告如下："], None),
        ("一、事实经过", ["__________"], None),
        ("二、法律意见", ["根据《__________》第____条等相关规定，____________________。"], None),
        ("三、本所要求", ["请贵方于 __________ 日前，____________________；逾期未予处理的，本所将建议委托人依法采取进一步措施。"], None),
     ],
     ["__________律师事务所（盖章）：__________",
      "经办律师：__________",
      "发函日期：__________"]),

    ("settlement-offer-cn.docx", "和解要约（不可作为证据）",
     [
        ("",
         ["致：__________", "（标注：本要约系不损害权利为目的，不得作为证据使用。）"], None),
        ("一、争议事项",
         ["双方争议事项摘要：____________________"], None),
        ("二、和解方案",
         ["方案要点：____________________"], None),
        ("三、答复期限",
         ["请贵方于 __________ 日前书面答复。逾期视为拒绝。"], None),
     ],
     ["发函方：__________", "发函日期：__________"]),
]


def main():
    cfg = json.loads(STYLE_CONFIG_PATH.read_text(encoding="utf-8"))
    TEMPLATE_DIR.mkdir(exist_ok=True)
    for filename, doc_title, body_sections, signature_lines in TEMPLATES:
        doc = _build(doc_title, body_sections, signature_lines, cfg)
        out = TEMPLATE_DIR / filename
        doc.save(out)
        print(f"[generate-templates] wrote {out.name}")
    print(f"[generate-templates] {len(TEMPLATES)} templates written.")


if __name__ == "__main__":
    main()
